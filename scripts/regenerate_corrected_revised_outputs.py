#!/usr/bin/env python3
"""Regenerate corrected revised figures and DOCX outputs.

Corrections handled here:
1. Regenerate Figure 1D from the corrected 3D pipeline, preserving zero-padded
   secondary-medical-area codes so Eastern Japan is not dropped during merges.
2. Regenerate Figure 2B as L008+L003 combined SCR, not L003/L008 ratio and not
   L008+L002.
3. Rebuild the revised English and Japanese DOCX files from the corrected figures.

This script is designed for GitHub Actions. It intentionally uses the legacy
/home/ubuntu paths expected by the existing generation scripts, while copying the
final products back into the repository tree.
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Iterable, Sequence

import geopandas as gpd
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from matplotlib.colors import BoundaryNorm
from matplotlib.lines import Line2D
from matplotlib.patches import Patch
from PIL import Image, ImageDraw, ImageFont, ImageOps

REPO_ROOT = Path(__file__).resolve().parents[1]
HOME_UBUNTU = Path("/home/ubuntu")
GIS_HOME = HOME_UBUNTU / "gis_data"
REPORT_PATH = REPO_ROOT / "reports" / "revised_figure_regeneration_report.md"

EN_MAP_DIR = REPO_ROOT / "output" / "maps_2d_en"
JP_MAP_DIR = REPO_ROOT / "output" / "maps_2d_jp"
OUTPUT_DIR = REPO_ROOT / "output"
REVISED_DOC_DIR = REPO_ROOT / "documents" / "revised"

ALL_BOUNDS = {"x": (122.5, 149.2), "y": (24.0, 46.0)}


class RegenerationError(RuntimeError):
    """Raised when an expected source or generated artifact is missing."""


def log(message: str) -> None:
    print(message, flush=True)


def run_command(command: Sequence[str], *, cwd: Path | None = None) -> None:
    """Run a command and fail loudly if it exits non-zero."""
    log("$ " + " ".join(command))
    subprocess.run(command, cwd=str(cwd or REPO_ROOT), check=True)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def copy_file(src: Path, dst: Path, *, required: bool = True) -> bool:
    """Copy a file, returning True when copied."""
    if not src.exists():
        if required:
            raise RegenerationError(f"Required file not found: {src}")
        log(f"Optional file not found, skipping: {src}")
        return False
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)
    log(f"Copied {src} -> {dst}")
    return True


def prepare_legacy_home_paths() -> None:
    """Prepare /home/ubuntu paths expected by existing scripts."""
    ensure_dir(HOME_UBUNTU)
    ensure_dir(REPORT_PATH.parent)

    source_gis = REPO_ROOT / "gis_data"
    if not source_gis.exists():
        raise RegenerationError(f"GIS directory not found: {source_gis}")

    if GIS_HOME.exists() or GIS_HOME.is_symlink():
        if GIS_HOME.is_symlink() or GIS_HOME.is_file():
            GIS_HOME.unlink()
        else:
            shutil.rmtree(GIS_HOME)
    GIS_HOME.symlink_to(source_gis, target_is_directory=True)
    log(f"Symlinked {GIS_HOME} -> {source_gis}")

    data_file_pairs = [
        (REPO_ROOT / "data" / "corrected_metrics_final.csv", HOME_UBUNTU / "corrected_metrics_final.csv"),
        (REPO_ROOT / "data" / "anesthesiologist_by_sma.csv", HOME_UBUNTU / "anesthesiologist_by_sma.csv"),
        (REPO_ROOT / "data" / "scr_n_kubun.csv", HOME_UBUNTU / "scr_n_kubun.csv"),
        (REPO_ROOT / "data" / "univ_hospital_mapping_v2.json", HOME_UBUNTU / "univ_hospital_mapping_v2.json"),
    ]
    for src, dst in data_file_pairs:
        copy_file(src, dst)

    # Existing revised DOCX scripts read composite figures from /home/ubuntu.
    for fig in OUTPUT_DIR.glob("rapm_fig*.png"):
        copy_file(fig, HOME_UBUNTU / fig.name, required=False)

    ensure_dir(HOME_UBUNTU / "3d_extruded")


def pick_font(size: int, *, bold: bool = False, japanese: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Pick a font available on GitHub-hosted Ubuntu runners."""
    candidates: list[str]
    if japanese:
        candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSerifCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont) -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x0, y0, x1, y1 = box
    x = x0 + (x1 - x0 - tw) // 2
    y = y0 + (y1 - y0 - th) // 2
    draw.text((x, y), text, fill=(0, 0, 0), font=font)


def compose_panel_figure(
    panels: Sequence[tuple[Path, str]],
    output_path: Path,
    *,
    columns: int,
    panel_size: tuple[int, int] = (1650, 1300),
    title_height: int = 100,
    margin: int = 80,
    gutter: int = 60,
    japanese: bool = False,
) -> None:
    """Compose a multi-panel figure from existing images."""
    if not panels:
        raise RegenerationError("No panels supplied for composite figure")
    for image_path, _ in panels:
        if not image_path.exists():
            raise RegenerationError(f"Panel image not found: {image_path}")

    rows = int(np.ceil(len(panels) / columns))
    panel_w, panel_h = panel_size
    total_w = margin * 2 + columns * panel_w + (columns - 1) * gutter
    total_h = margin * 2 + rows * panel_h + (rows - 1) * gutter
    canvas = Image.new("RGB", (total_w, total_h), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pick_font(42, bold=True, japanese=japanese)

    for idx, (image_path, title) in enumerate(panels):
        row = idx // columns
        col = idx % columns
        x = margin + col * (panel_w + gutter)
        y = margin + row * (panel_h + gutter)
        title_box = (x, y, x + panel_w, y + title_height)
        draw_centered_text(draw, title_box, title, title_font)

        image_box = (x, y + title_height, x + panel_w, y + panel_h)
        target_w = image_box[2] - image_box[0]
        target_h = image_box[3] - image_box[1]
        with Image.open(image_path) as im:
            im = im.convert("RGB")
            contained = ImageOps.contain(im, (target_w, target_h), method=Image.Resampling.LANCZOS)
        paste_x = image_box[0] + (target_w - contained.width) // 2
        paste_y = image_box[1] + (target_h - contained.height) // 2
        canvas.paste(contained, (paste_x, paste_y))

        # Thin panel border for visual separation.
        draw.rectangle((x, y, x + panel_w, y + panel_h), outline=(180, 180, 180), width=2)

    ensure_dir(output_path.parent)
    canvas.save(output_path, dpi=(300, 300), quality=95)
    log(f"Saved composite figure: {output_path}")


def make_l008_l003_combined_maps() -> dict[str, int | float | str]:
    """Create corrected L008+L003 combined SCR maps in English and Japanese."""
    gdf = gpd.read_file(GIS_HOME / "merged_enriched.gpkg")
    gdf_pref = gpd.read_file(GIS_HOME / "pref_simplified.gpkg")
    gdf_nt = gpd.read_file(GIS_HOME / "northern_territories.gpkg")

    for required_col in ["L008_scr", "L003_scr"]:
        if required_col not in gdf.columns:
            raise RegenerationError(f"Column {required_col} not present in merged_enriched.gpkg")

    gdf["L008_L003_combined"] = gdf["L008_scr"] + gdf["L003_scr"]
    valid = gdf["L008_L003_combined"].notna()

    boundaries = [0, 50, 100, 150, 200, 250, 300, 400, 550, 800, 1000]

    def make_map(title: str, legend_label: str, filename: Path, *, japanese: bool) -> None:
        if japanese:
            try:
                import japanize_matplotlib  # noqa: F401
            except Exception:
                plt.rcParams["font.family"] = "Noto Sans CJK JP"
        fig, ax = plt.subplots(1, 1, figsize=(14, 18))
        norm = BoundaryNorm(boundaries, ncolors=256)
        nodata = gdf[gdf["L008_L003_combined"].isna()]
        valid_gdf = gdf[gdf["L008_L003_combined"].notna()]
        if len(nodata) > 0:
            nodata.plot(ax=ax, color="#e0e0e0", edgecolor="none")
        valid_gdf.plot(
            ax=ax,
            column="L008_L003_combined",
            cmap="RdYlBu_r",
            norm=norm,
            edgecolor="none",
            legend=False,
        )
        gdf_nt.plot(ax=ax, color="white", edgecolor="#333333", linewidth=0.5)
        gdf.boundary.plot(ax=ax, linewidth=0.3, color="#888888", linestyle=(0, (2, 3)))
        gdf_pref.boundary.plot(ax=ax, linewidth=0.8, color="#333333", linestyle="solid")
        if "has_univ" in gdf.columns:
            univ = gdf[gdf["has_univ"] == 1]
            if len(univ) > 0:
                centroids = univ.geometry.centroid
                ax.scatter(
                    centroids.x,
                    centroids.y,
                    s=18,
                    c="red",
                    marker="o",
                    zorder=5,
                    linewidths=0.5,
                    edgecolors="darkred",
                    alpha=0.8,
                )
        ax.set_xlim(ALL_BOUNDS["x"])
        ax.set_ylim(ALL_BOUNDS["y"])
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(title, fontsize=15, fontweight="bold", pad=15)

        if japanese:
            line_items = [
                Line2D([0], [0], color="#333333", lw=1.0, ls="solid", label="都道府県境"),
                Line2D([0], [0], color="#888888", lw=0.5, ls=(0, (2, 3)), label="二次医療圏境"),
                Patch(facecolor="white", edgecolor="#333333", linewidth=0.5, label="北方領土（医療圏未設定）"),
                Line2D([0], [0], marker="o", color="w", markerfacecolor="red", markersize=6, label="大学病院所在圏"),
            ]
        else:
            line_items = [
                Line2D([0], [0], color="#333333", lw=1.0, ls="solid", label="Prefecture boundary"),
                Line2D([0], [0], color="#888888", lw=0.5, ls=(0, (2, 3)), label="SMA boundary"),
                Patch(facecolor="white", edgecolor="#333333", linewidth=0.5, label="Northern Territories (no SMA)"),
                Line2D([0], [0], marker="o", color="w", markerfacecolor="red", markersize=6, label="University hospital SMA"),
            ]
        ax.legend(handles=line_items, loc="lower left", fontsize=8.5, framealpha=0.95, edgecolor="gray")

        sm = plt.cm.ScalarMappable(cmap="RdYlBu_r", norm=norm)
        sm.set_array([])
        cbar_ax = fig.add_axes([0.15, 0.06, 0.55, 0.015])
        cbar = fig.colorbar(sm, cax=cbar_ax, orientation="horizontal")
        cbar.set_label(legend_label, fontsize=12)

        ensure_dir(filename.parent)
        plt.savefig(filename, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        log(f"Saved corrected L008+L003 map: {filename}")

    en_path = HOME_UBUNTU / "en_map_L008_L003_combined.png"
    jp_path = HOME_UBUNTU / "map_L008_L003_combined.png"
    make_map(
        "Combined General Anesthesia + Continuous Epidural Infusion SCR (L008+L003)",
        "Combined SCR",
        en_path,
        japanese=False,
    )
    make_map(
        "全身麻酔+持続硬膜外 合計SCR (L008+L003)",
        "合計SCR",
        jp_path,
        japanese=True,
    )

    copy_file(en_path, EN_MAP_DIR / "en_map_L008_L003_combined.png")
    copy_file(jp_path, JP_MAP_DIR / "map_L008_L003_combined.png")

    stats = {
        "valid_l008_l003_smas": int(valid.sum()),
        "total_smas": int(len(gdf)),
        "mean_l008_l003_combined": float(gdf.loc[valid, "L008_L003_combined"].mean()),
        "median_l008_l003_combined": float(gdf.loc[valid, "L008_L003_combined"].median()),
        "max_l008_l003_combined": float(gdf.loc[valid, "L008_L003_combined"].max()),
    }
    return stats


def copy_generated_standard_maps_to_repo() -> None:
    """Copy maps generated by existing scripts from /home/ubuntu back to output directories."""
    en_names = [
        "en_map_L008_scr.png",
        "en_map_univ_presence.png",
        "en_map_L004_scr.png",
        "en_map_L008_L004_combined.png",
        "en_map_L003_scr.png",
        "en_map_L003_L008_ratio_corrected.png",
    ]
    jp_names = [
        "map_L008_scr.png",
        "map_univ_presence.png",
        "map_n_univ.png",
        "map_L004_scr.png",
        "map_L008_L004_combined.png",
        "map_L002_scr.png",
        "map_L003_scr.png",
        "map_L002_L008_ratio.png",
        "map_L003_L008_ratio.png",
        "map_L008_L002_combined.png",
    ]
    for name in en_names:
        copy_file(HOME_UBUNTU / name, EN_MAP_DIR / name, required=False)
    for name in jp_names:
        copy_file(HOME_UBUNTU / name, JP_MAP_DIR / name, required=False)


def copy_3d_outputs_to_repo() -> None:
    target_dir = REPO_ROOT / "output" / "maps_3d_en"
    ensure_dir(target_dir)
    for name in ["3D_ratio_by_anes_v2.png", "3D_L008_by_anes_v2.png", "3D_ratio_by_surgery_v2.png"]:
        copy_file(HOME_UBUNTU / "3d_extruded" / name, target_dir / name, required=False)


def create_corrected_composite_figures() -> None:
    """Create corrected RAPM composite Figure 1 and Figure 2 images."""
    fig1_en_panels = [
        (HOME_UBUNTU / "en_map_L008_scr.png", "(A) General anesthesia (L008) SCR"),
        (HOME_UBUNTU / "en_map_L004_scr.png", "(B) Spinal anesthesia (L004) SCR"),
        (HOME_UBUNTU / "en_map_L008_L004_combined.png", "(C) L008+L004 combined SCR"),
        (HOME_UBUNTU / "3d_extruded" / "3D_L008_by_anes_v2.png", "(D) 3D L008 SCR x anesthesiologist count"),
    ]
    fig1_jp_panels = [
        (HOME_UBUNTU / "map_L008_scr.png", "(A) 全身麻酔 (L008) SCR"),
        (HOME_UBUNTU / "map_L004_scr.png", "(B) 脊椎麻酔 (L004) SCR"),
        (HOME_UBUNTU / "map_L008_L004_combined.png", "(C) L008+L004 合計SCR"),
        (HOME_UBUNTU / "3d_extruded" / "3D_L008_by_anes_v2.png", "(D) 3D: L008 SCR x 麻酔科医数"),
    ]
    fig2_en_panels = [
        (HOME_UBUNTU / "en_map_L003_scr.png", "(A) Continuous epidural infusion (L003) SCR"),
        (HOME_UBUNTU / "en_map_L008_L003_combined.png", "(B) L008+L003 combined SCR"),
    ]
    fig2_jp_panels = [
        (HOME_UBUNTU / "map_L003_scr.png", "(A) 硬膜外麻酔後持続注入 (L003) SCR"),
        (HOME_UBUNTU / "map_L008_L003_combined.png", "(B) L008+L003 合計SCR"),
    ]

    compose_panel_figure(fig1_en_panels, OUTPUT_DIR / "rapm_fig1_en.png", columns=2, japanese=False)
    compose_panel_figure(fig1_jp_panels, OUTPUT_DIR / "rapm_fig1_jp.png", columns=2, japanese=True)
    compose_panel_figure(
        fig2_en_panels,
        OUTPUT_DIR / "rapm_fig2_en.png",
        columns=2,
        panel_size=(1650, 1350),
        japanese=False,
    )
    compose_panel_figure(
        fig2_jp_panels,
        OUTPUT_DIR / "rapm_fig2_jp.png",
        columns=2,
        panel_size=(1650, 1350),
        japanese=True,
    )

    # The revised manuscript scripts read figures from /home/ubuntu.
    for name in ["rapm_fig1_en.png", "rapm_fig1_jp.png", "rapm_fig2_en.png", "rapm_fig2_jp.png"]:
        copy_file(OUTPUT_DIR / name, HOME_UBUNTU / name)


def patch_caption_paragraph(docx_path: Path, *, language: str) -> dict[str, bool]:
    """Patch Figure 1 and Figure 2 captions in a generated DOCX."""
    if not docx_path.exists():
        raise RegenerationError(f"DOCX to patch not found: {docx_path}")
    doc = Document(str(docx_path))
    found = {"figure1": False, "figure2": False}

    if language == "en":
        fig1_caption = (
            "Figure 1. Geographic distribution of anesthesia-related standardized claim ratios. "
            "(A) General anesthesia (L008) SCR. (B) Spinal anesthesia (L004) SCR. "
            "(C) Combined L008+L004 SCR. (D) Three-dimensional extruded map of L008 SCR "
            "with anesthesiologist count as height; regenerated with zero-padded SMA codes so "
            "Eastern Japan is retained in the data merge."
        )
        fig2_caption = (
            "Figure 2. Regional distribution of combined general-epidural anesthesia indicators. "
            "(A) Continuous epidural infusion (L003) SCR. "
            "(B) Combined L008+L003 SCR, representing general anesthesia plus continuous epidural "
            "infusion claim burden. Panel B intentionally uses L008+L003, not L003/L008."
        )
        fig1_prefixes = ("Figure 1",)
        fig2_prefixes = ("Figure 2",)
    else:
        fig1_caption = (
            "図1. 麻酔関連SCRの地理的分布。"
            "(A) 全身麻酔（L008）SCR。(B) 脊椎麻酔（L004）SCR。"
            "(C) L008+L004合計SCR。(D) L008 SCRを色、麻酔科医数を高さで示す3D押し出し図。"
            "二次医療圏コードをゼロ埋めして再生成し、東日本のデータが欠落しないように修正した。"
        )
        fig2_caption = (
            "図2. 全身麻酔+硬膜外併用指標の地理的分布。"
            "(A) 硬膜外麻酔後持続注入（L003）SCR。"
            "(B) L008+L003合計SCR。パネルBはL003/L008比ではなく、L008+L003として再生成した。"
        )
        fig1_prefixes = ("図1", "Figure 1")
        fig2_prefixes = ("図2", "Figure 2")

    def replace_text(paragraph, new_text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped.startswith(fig1_prefixes):
            replace_text(paragraph, fig1_caption)
            found["figure1"] = True
        elif stripped.startswith(fig2_prefixes):
            replace_text(paragraph, fig2_caption)
            found["figure2"] = True

    doc.save(str(docx_path))
    return found


def regenerate_docx_files() -> dict[str, dict[str, bool]]:
    """Run existing revised manuscript scripts and patch captions."""
    run_command([sys.executable, "scripts/create_revised_rapm_en.py"])
    run_command([sys.executable, "scripts/create_revised_rapm_jp.py"])

    en_src = HOME_UBUNTU / "regional_anesthesia_REVISED_EN.docx"
    jp_src = HOME_UBUNTU / "regional_anesthesia_REVISED_JP.docx"
    en_dst = REVISED_DOC_DIR / "regional_anesthesia_REVISED_EN.docx"
    jp_dst = REVISED_DOC_DIR / "regional_anesthesia_REVISED_JP.docx"
    copy_file(en_src, en_dst)
    copy_file(jp_src, jp_dst)

    found = {
        "en": patch_caption_paragraph(en_dst, language="en"),
        "jp": patch_caption_paragraph(jp_dst, language="jp"),
    }
    return found


def verify_zero_padded_eastern_merge() -> dict[str, int]:
    """Verify that East Japan area codes retain anesthesiologist merge coverage."""
    gdf = gpd.read_file(GIS_HOME / "merged_enriched.gpkg")
    anes = pd.read_csv(HOME_UBUNTU / "anesthesiologist_by_sma.csv")
    if "code" in anes.columns and "area_code" not in anes.columns:
        anes = anes.rename(columns={"code": "area_code"})
    if "anesthesiologist_count" in anes.columns and "n_anes" not in anes.columns:
        anes = anes.rename(columns={"anesthesiologist_count": "n_anes"})
    if "area_code" not in gdf.columns or "area_code" not in anes.columns:
        return {"east_total": -1, "east_with_anesthesiologist_count": -1}
    gdf = gdf.copy()
    anes = anes.copy()
    gdf["area_code"] = gdf["area_code"].astype(str).str.zfill(4)
    anes["area_code"] = anes["area_code"].astype(str).str.zfill(4)
    merged = gdf[["area_code"]].merge(anes[["area_code", "n_anes"]], on="area_code", how="left")
    east_prefixes = {f"{i:02d}" for i in range(1, 15)}  # Hokkaido through Kanagawa.
    east = merged[merged["area_code"].str[:2].isin(east_prefixes)]
    return {
        "east_total": int(len(east)),
        "east_with_anesthesiologist_count": int(east["n_anes"].notna().sum()),
    }


def write_report(
    combined_stats: dict[str, int | float | str],
    east_stats: dict[str, int],
    caption_patch_status: dict[str, dict[str, bool]],
) -> None:
    """Write a human-readable regeneration report."""
    lines = [
        "# Revised figure/DOCX regeneration report",
        "",
        "This report was generated by `scripts/regenerate_corrected_revised_outputs.py`.",
        "",
        "## Corrections applied",
        "",
        "1. **Figure 1D** was regenerated from the corrected 3D pipeline. The 3D pipeline zero-pads `area_code` before joining auxiliary data, preventing Hokkaido/Tohoku/Kanto/Eastern Japan rows from being dropped during merges.",
        "2. **Figure 2B** was regenerated as **L008+L003 combined SCR**. It is no longer the L003/L008 ratio and is not L008+L002.",
        "3. Revised English and Japanese DOCX manuscripts were rebuilt and their Figure 1/Figure 2 captions were patched to match the corrected panels.",
        "",
        "## Verification counts",
        "",
        f"- L008+L003 valid SMAs: {combined_stats['valid_l008_l003_smas']} / {combined_stats['total_smas']}",
        f"- L008+L003 mean SCR: {combined_stats['mean_l008_l003_combined']:.3f}",
        f"- L008+L003 median SCR: {combined_stats['median_l008_l003_combined']:.3f}",
        f"- L008+L003 max SCR: {combined_stats['max_l008_l003_combined']:.3f}",
        f"- Eastern Japan SMAs with anesthesiologist-count merge coverage: {east_stats['east_with_anesthesiologist_count']} / {east_stats['east_total']}",
        "",
        "## Caption patch status",
        "",
        "```json",
        json.dumps(caption_patch_status, ensure_ascii=False, indent=2),
        "```",
        "",
        "## Generated files",
        "",
        "- `output/rapm_fig1_en.png`",
        "- `output/rapm_fig1_jp.png`",
        "- `output/rapm_fig2_en.png`",
        "- `output/rapm_fig2_jp.png`",
        "- `output/maps_2d_en/en_map_L008_L003_combined.png`",
        "- `output/maps_2d_jp/map_L008_L003_combined.png`",
        "- `documents/revised/regional_anesthesia_REVISED_EN.docx`",
        "- `documents/revised/regional_anesthesia_REVISED_JP.docx`",
        "",
    ]
    ensure_dir(REPORT_PATH.parent)
    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")
    log(f"Saved report: {REPORT_PATH}")


def main() -> None:
    os.chdir(REPO_ROOT)
    prepare_legacy_home_paths()

    # Rebuild base maps and 3D screenshots using the repository's existing scripts.
    run_command([sys.executable, "scripts/generate_maps_en.py"])
    run_command([sys.executable, "scripts/generate_all_maps.py"])
    combined_stats = make_l008_l003_combined_maps()
    copy_generated_standard_maps_to_repo()

    run_command([sys.executable, "scripts/create_3d_v2.py"])
    run_command([sys.executable, "scripts/screenshot_3d_v2.py"])
    copy_3d_outputs_to_repo()
    east_stats = verify_zero_padded_eastern_merge()

    create_corrected_composite_figures()
    caption_patch_status = regenerate_docx_files()
    write_report(combined_stats, east_stats, caption_patch_status)

    log("Regeneration completed successfully.")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        log(f"ERROR: {exc}")
        raise
